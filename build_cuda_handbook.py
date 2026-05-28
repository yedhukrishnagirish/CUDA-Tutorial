from pathlib import Path
from textwrap import dedent
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as PdfImage,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "CUDA_Cpp_AI_ML_Interview_Handbook.docx"
PDF_OUT = ROOT / "CUDA_Cpp_AI_ML_Interview_Handbook.pdf"
SAMPLES = ROOT / "cuda_samples"
FIGS = ROOT / "figures"
SAMPLES.mkdir(exist_ok=True)
FIGS.mkdir(exist_ok=True)


PALETTE = {
    "ink": "172033",
    "blue": "245B91",
    "teal": "1F7A72",
    "gold": "B7791F",
    "red": "A23B3B",
    "green": "3A7D44",
    "gray": "F2F4F7",
    "line": "CBD5E1",
    "dark": "0F172A",
}


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_cell(cell, text, bold=False, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(9)
    r.font.name = "Calibri"
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell(hdr[i], h, bold=True, fill="E8EEF5")
        if widths:
            hdr[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell(cells[i], val)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_code(doc, code):
    p = doc.add_paragraph()
    for line in dedent(code).strip("\n").splitlines():
        r = p.add_run(line.rstrip() + "\n")
        r.font.name = "Consolas"
        r.font.size = Pt(8)
        r.font.color.rgb = rgb("111827")
    p.paragraph_format.space_after = Pt(6)


def draw_box(draw, xy, text, fill, outline="#334155", font=None):
    draw.rounded_rectangle(xy, radius=14, fill=fill, outline=outline, width=2)
    x1, y1, x2, y2 = xy
    lines = text.split("\n")
    line_h = 20
    total_h = len(lines) * line_h
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, fill="#0F172A", font=font)
        y += line_h


def make_figures():
    font = ImageFont.truetype("arial.ttf", 16)
    small = ImageFont.truetype("arial.ttf", 13)
    title = ImageFont.truetype("arial.ttf", 22)

    # Execution hierarchy
    img = Image.new("RGB", (1200, 520), "white")
    d = ImageDraw.Draw(img)
    d.text((40, 24), "CUDA Execution Hierarchy", fill="#0F172A", font=title)
    draw_box(d, (60, 90, 1140, 165), "Grid = one kernel launch", "#DBEAFE", font=font)
    for i, x in enumerate([120, 380, 640, 900]):
        draw_box(d, (x, 220, x + 170, 310), f"Block {i}\nshared memory", "#DCFCE7", font=font)
        for j in range(4):
            draw_box(d, (x + j * 38, 360, x + 30 + j * 38, 410), f"T{j}", "#FEF3C7", font=small)
        d.line((x + 85, 165, x + 85, 220), fill="#64748B", width=3)
        d.line((x + 85, 310, x + 85, 360), fill="#64748B", width=3)
    d.text((70, 455), "A kernel launches a grid. A grid has blocks. Blocks have threads. Threads in a block can cooperate through shared memory and barriers.", fill="#334155", font=small)
    img.save(FIGS / "execution_hierarchy.png")

    # Memory hierarchy
    img = Image.new("RGB", (1200, 580), "white")
    d = ImageDraw.Draw(img)
    d.text((40, 24), "CUDA Memory Hierarchy", fill="#0F172A", font=title)
    layers = [
        ((140, 95, 1060, 160), "Registers: per-thread, fastest, limited", "#FEF3C7"),
        ((190, 185, 1010, 250), "Shared memory / L1: per-block, explicit, low latency", "#DCFCE7"),
        ((240, 275, 960, 340), "L2 cache: device-wide, shared by SMs", "#DBEAFE"),
        ((290, 365, 910, 430), "Global memory: high bandwidth, high latency", "#E0E7FF"),
        ((340, 455, 860, 520), "Host memory: CPU RAM, reached through PCIe/NVLink", "#F3E8FF"),
    ]
    for xy, text, fill in layers:
        draw_box(d, xy, text, fill, font=font)
    d.text((60, 540), "Interview rule: performance usually starts with reducing global-memory traffic and making remaining accesses coalesced.", fill="#334155", font=small)
    img.save(FIGS / "memory_hierarchy.png")

    # SM architecture
    img = Image.new("RGB", (1200, 650), "white")
    d = ImageDraw.Draw(img)
    d.text((40, 24), "Streaming Multiprocessor Mental Model", fill="#0F172A", font=title)
    draw_box(d, (80, 110, 1120, 580), "SM", "#F8FAFC", font=title)
    blocks = [
        ((130, 170, 410, 250), "Warp schedulers\nissue ready warps", "#DBEAFE"),
        ((460, 170, 740, 250), "CUDA cores\nFP32 / INT work", "#DCFCE7"),
        ((790, 170, 1070, 250), "Tensor cores\nmatrix math", "#FEF3C7"),
        ((130, 310, 410, 390), "Registers\nper active thread", "#FDE68A"),
        ((460, 310, 740, 390), "Shared memory / L1\nper block cooperation", "#BBF7D0"),
        ((790, 310, 1070, 390), "Load/store units\nmemory instructions", "#BFDBFE"),
        ((300, 455, 900, 535), "Many resident warps hide latency by switching when one warp waits", "#E9D5FF"),
    ]
    for xy, text, fill in blocks:
        draw_box(d, xy, text, fill, font=font)
    img.save(FIGS / "sm_model.png")

    # Optimization loop
    img = Image.new("RGB", (1200, 500), "white")
    d = ImageDraw.Draw(img)
    d.text((40, 24), "CUDA Optimization Loop", fill="#0F172A", font=title)
    steps = [
        ("Correct kernel", "#DBEAFE"),
        ("Measure", "#DCFCE7"),
        ("Find bottleneck", "#FEF3C7"),
        ("Change one thing", "#FCE7F3"),
        ("Re-measure", "#E0E7FF"),
    ]
    xs = [70, 295, 520, 745, 970]
    for i, (text, fill) in enumerate(steps):
        draw_box(d, (xs[i], 170, xs[i] + 160, 260), text, fill, font=font)
        if i < len(steps) - 1:
            d.line((xs[i] + 160, 215, xs[i + 1], 215), fill="#64748B", width=4)
            d.polygon([(xs[i + 1], 215), (xs[i + 1] - 14, 207), (xs[i + 1] - 14, 223)], fill="#64748B")
    d.text((80, 340), "Never optimize from vibes. Use CUDA events for timing, Nsight Systems for timeline, and Nsight Compute for kernel-level evidence.", fill="#334155", font=font)
    img.save(FIGS / "optimization_loop.png")


SAMPLE_FILES = {
    "01_hello.cu": r'''
        #include <cstdio>

        __global__ void hello() {
            printf("Hello from block %d, thread %d\n", blockIdx.x, threadIdx.x);
        }

        int main() {
            hello<<<2, 4>>>();
            cudaDeviceSynchronize();
            return 0;
        }
    ''',
    "02_vector_add.cu": r'''
        #include <cuda_runtime.h>
        #include <iostream>
        #include <vector>

        #define CUDA_CHECK(call) do {                                      \
            cudaError_t err = call;                                        \
            if (err != cudaSuccess) {                                      \
                std::cerr << "CUDA error: " << cudaGetErrorString(err)     \
                          << " at " << __FILE__ << ":" << __LINE__ << "\n"; \
                std::exit(1);                                              \
            }                                                              \
        } while (0)

        __global__ void vector_add(const float* a, const float* b, float* c, int n) {
            int i = blockIdx.x * blockDim.x + threadIdx.x;
            if (i < n) c[i] = a[i] + b[i];
        }

        int main() {
            int n = 1 << 20;
            size_t bytes = n * sizeof(float);
            std::vector<float> h_a(n, 1.0f), h_b(n, 2.0f), h_c(n);
            float *d_a = nullptr, *d_b = nullptr, *d_c = nullptr;

            CUDA_CHECK(cudaMalloc(&d_a, bytes));
            CUDA_CHECK(cudaMalloc(&d_b, bytes));
            CUDA_CHECK(cudaMalloc(&d_c, bytes));
            CUDA_CHECK(cudaMemcpy(d_a, h_a.data(), bytes, cudaMemcpyHostToDevice));
            CUDA_CHECK(cudaMemcpy(d_b, h_b.data(), bytes, cudaMemcpyHostToDevice));

            int threads = 256;
            int blocks = (n + threads - 1) / threads;
            vector_add<<<blocks, threads>>>(d_a, d_b, d_c, n);
            CUDA_CHECK(cudaGetLastError());
            CUDA_CHECK(cudaMemcpy(h_c.data(), d_c, bytes, cudaMemcpyDeviceToHost));

            std::cout << "c[123] = " << h_c[123] << "\n";
            cudaFree(d_a); cudaFree(d_b); cudaFree(d_c);
            return 0;
        }
    ''',
    "03_matrix_mul_tiled.cu": r'''
        #include <cuda_runtime.h>
        #include <iostream>
        #include <vector>

        constexpr int TILE = 16;

        __global__ void matmul_tiled(const float* A, const float* B, float* C, int N) {
            __shared__ float As[TILE][TILE];
            __shared__ float Bs[TILE][TILE];

            int row = blockIdx.y * TILE + threadIdx.y;
            int col = blockIdx.x * TILE + threadIdx.x;
            float sum = 0.0f;

            for (int t = 0; t < (N + TILE - 1) / TILE; ++t) {
                int a_col = t * TILE + threadIdx.x;
                int b_row = t * TILE + threadIdx.y;
                As[threadIdx.y][threadIdx.x] = (row < N && a_col < N) ? A[row * N + a_col] : 0.0f;
                Bs[threadIdx.y][threadIdx.x] = (b_row < N && col < N) ? B[b_row * N + col] : 0.0f;
                __syncthreads();

                for (int k = 0; k < TILE; ++k) {
                    sum += As[threadIdx.y][k] * Bs[k][threadIdx.x];
                }
                __syncthreads();
            }
            if (row < N && col < N) C[row * N + col] = sum;
        }

        int main() {
            int N = 512;
            size_t bytes = N * N * sizeof(float);
            std::vector<float> A(N * N, 1.0f), B(N * N, 2.0f), C(N * N);
            float *d_A, *d_B, *d_C;
            cudaMalloc(&d_A, bytes); cudaMalloc(&d_B, bytes); cudaMalloc(&d_C, bytes);
            cudaMemcpy(d_A, A.data(), bytes, cudaMemcpyHostToDevice);
            cudaMemcpy(d_B, B.data(), bytes, cudaMemcpyHostToDevice);
            dim3 threads(TILE, TILE);
            dim3 blocks((N + TILE - 1) / TILE, (N + TILE - 1) / TILE);
            matmul_tiled<<<blocks, threads>>>(d_A, d_B, d_C, N);
            cudaMemcpy(C.data(), d_C, bytes, cudaMemcpyDeviceToHost);
            std::cout << "C[0] = " << C[0] << "\n";
            cudaFree(d_A); cudaFree(d_B); cudaFree(d_C);
            return 0;
        }
    ''',
    "04_reduction.cu": r'''
        #include <cuda_runtime.h>
        #include <iostream>
        #include <numeric>
        #include <vector>

        __global__ void reduce_sum(const float* in, float* block_sums, int n) {
            extern __shared__ float s[];
            unsigned int tid = threadIdx.x;
            unsigned int i = blockIdx.x * blockDim.x * 2 + threadIdx.x;
            float x = 0.0f;
            if (i < n) x += in[i];
            if (i + blockDim.x < n) x += in[i + blockDim.x];
            s[tid] = x;
            __syncthreads();

            for (unsigned int stride = blockDim.x / 2; stride > 0; stride >>= 1) {
                if (tid < stride) s[tid] += s[tid + stride];
                __syncthreads();
            }
            if (tid == 0) block_sums[blockIdx.x] = s[0];
        }

        int main() {
            int n = 1 << 20;
            std::vector<float> h(n, 1.0f);
            int threads = 256;
            int blocks = (n + threads * 2 - 1) / (threads * 2);
            float *d_in, *d_sums;
            cudaMalloc(&d_in, n * sizeof(float));
            cudaMalloc(&d_sums, blocks * sizeof(float));
            cudaMemcpy(d_in, h.data(), n * sizeof(float), cudaMemcpyHostToDevice);
            reduce_sum<<<blocks, threads, threads * sizeof(float)>>>(d_in, d_sums, n);
            std::vector<float> partial(blocks);
            cudaMemcpy(partial.data(), d_sums, blocks * sizeof(float), cudaMemcpyDeviceToHost);
            float total = std::accumulate(partial.begin(), partial.end(), 0.0f);
            std::cout << "sum = " << total << "\n";
            cudaFree(d_in); cudaFree(d_sums);
            return 0;
        }
    ''',
    "05_streams_async.cu": r'''
        #include <cuda_runtime.h>
        #include <iostream>

        __global__ void scale(float* x, int n, float a) {
            int i = blockIdx.x * blockDim.x + threadIdx.x;
            if (i < n) x[i] *= a;
        }

        int main() {
            int n = 1 << 20;
            size_t bytes = n * sizeof(float);
            float *h = nullptr, *d = nullptr;
            cudaMallocHost(&h, bytes);       // pinned host memory enables async copies
            cudaMalloc(&d, bytes);
            for (int i = 0; i < n; ++i) h[i] = 1.0f;

            cudaStream_t stream;
            cudaStreamCreate(&stream);
            cudaMemcpyAsync(d, h, bytes, cudaMemcpyHostToDevice, stream);
            scale<<<(n + 255) / 256, 256, 0, stream>>>(d, n, 3.0f);
            cudaMemcpyAsync(h, d, bytes, cudaMemcpyDeviceToHost, stream);
            cudaStreamSynchronize(stream);

            std::cout << "h[0] = " << h[0] << "\n";
            cudaStreamDestroy(stream);
            cudaFree(d);
            cudaFreeHost(h);
            return 0;
        }
    ''',
    "06_unified_memory.cu": r'''
        #include <cuda_runtime.h>
        #include <iostream>

        __global__ void add_one(float* x, int n) {
            int i = blockIdx.x * blockDim.x + threadIdx.x;
            if (i < n) x[i] += 1.0f;
        }

        int main() {
            int n = 1 << 20;
            float* x = nullptr;
            cudaMallocManaged(&x, n * sizeof(float));
            for (int i = 0; i < n; ++i) x[i] = 41.0f;

            add_one<<<(n + 255) / 256, 256>>>(x, n);
            cudaDeviceSynchronize();         // required before CPU reads managed data
            std::cout << "x[0] = " << x[0] << "\n";
            cudaFree(x);
            return 0;
        }
    ''',
}


def write_samples():
    for name, code in SAMPLE_FILES.items():
        (SAMPLES / name).write_text(dedent(code).strip() + "\n", encoding="utf-8")


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(PALETTE["ink"])
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Heading 1", 16, PALETTE["blue"], 14, 7),
        ("Heading 2", 13, PALETTE["teal"], 10, 5),
        ("Heading 3", 11.5, PALETTE["dark"], 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.bold = True
        st.font.size = Pt(size)
        st.font.color.rgb = rgb(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("CUDA C++ AI/ML Interview Handbook")
    return doc


def callout(doc, label, text, fill="F8FAFC"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell(cell, f"{label}: {text}", bold=False, fill=fill)
    doc.add_paragraph()


def build_doc():
    make_figures()
    write_samples()
    doc = setup_doc()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("CUDA C++ for AI/ML Interviews")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = rgb(PALETTE["blue"])
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Beginner to expert roadmap, architecture notes, runnable code, execution commands, and interview drills").italic = True
    doc.add_paragraph("Audience: a full-stack engineer moving into C++ CUDA, GPU computing, AI/ML systems, inference, training, and performance engineering.")
    callout(doc, "How to use this PDF", "Read it in passes: first understand the mental model, then compile every sample, then profile and explain each optimization aloud.", "E8EEF5")

    doc.add_heading("0. What CUDA C++ Is", level=1)
    doc.add_paragraph("CUDA C++ is C++ plus NVIDIA GPU extensions. You write normal host code that runs on the CPU and kernel code that runs on the GPU. The GPU is not a faster CPU; it is a throughput machine designed to run many lightweight threads and hide memory latency with massive concurrency.")
    add_bullets(doc, [
        "CPU code is called host code. GPU code is called device code.",
        "A kernel is a GPU function launched from the CPU with the <<<grid, block>>> syntax.",
        "A thread is the smallest CUDA execution unit you program directly.",
        "A warp is a hardware scheduling group, commonly 32 threads.",
        "A block is a group of threads that can synchronize and share fast shared memory.",
        "A grid is all blocks created by one kernel launch.",
    ])

    doc.add_picture(str(FIGS / "execution_hierarchy.png"), width=Inches(6.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. Setup And First Execution", level=1)
    add_numbers(doc, [
        "Install an NVIDIA driver that supports your GPU.",
        "Install the CUDA Toolkit. On Windows, install a compatible Visual Studio C++ toolchain first.",
        "Confirm the toolchain with: nvcc --version",
        "Confirm the GPU and driver with: nvidia-smi",
        "Compile a .cu file with: nvcc file.cu -o file.exe",
        "Run the executable from PowerShell: .\\file.exe",
    ])
    add_code(doc, r"""
        cd C:\Users\ygirish\Documents\Codex\2026-05-28\hi\cuda_samples
        nvcc 01_hello.cu -o 01_hello.exe
        .\01_hello.exe
    """)

    doc.add_heading("2. CUDA Language Keywords", level=1)
    add_table(doc, ["Keyword", "Meaning", "Interview signal"], [
        ["__global__", "Function runs on GPU and is launched by CPU.", "Used for kernels; returns void."],
        ["__device__", "Function or variable lives on GPU device.", "Callable from device code."],
        ["__host__", "Function runs on CPU host.", "Can combine with __device__ for dual-use helpers."],
        ["__shared__", "Per-block shared memory.", "Fast cooperation inside one block."],
        ["__syncthreads()", "Barrier for threads in a block.", "Needed before reading shared data written by peers."],
        ["threadIdx/blockIdx", "Built-in coordinates.", "Used to map work to data."],
        ["blockDim/gridDim", "Launch dimensions.", "Used for bounds and grid-stride loops."],
    ], widths=[1.35, 3.0, 2.1])

    doc.add_heading("3. Memory Model", level=1)
    doc.add_picture(str(FIGS / "memory_hierarchy.png"), width=Inches(6.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_table(doc, ["Memory", "Lifetime / scope", "Typical use", "Pitfall"], [
        ["Registers", "One thread", "Local scalars", "Too many registers can reduce occupancy."],
        ["Local memory", "One thread, backed by global", "Spilled arrays/variables", "Slow if compiler spills registers."],
        ["Shared memory", "One block", "Tiling, reductions, stencil reuse", "Bank conflicts and missing barriers."],
        ["Global memory", "Whole device allocation", "Large arrays and tensors", "Latency, non-coalesced access."],
        ["Constant memory", "Device-wide read-only", "Small broadcast constants", "Bad for random per-thread access."],
        ["Unified memory", "Managed CPU/GPU address", "Simplicity and prototyping", "Page migration surprises."],
        ["Pinned host memory", "CPU page-locked", "Async copies and overlap", "Overuse hurts system memory."],
    ], widths=[1.2, 1.4, 2.0, 1.9])
    callout(doc, "Interview answer", "Coalescing means adjacent threads in a warp access adjacent memory addresses, allowing the hardware to combine requests efficiently.", "FEF3C7")

    doc.add_heading("4. Writing Correct Kernels", level=1)
    doc.add_paragraph("The most common CUDA beginner bug is forgetting that you usually launch more threads than data elements. Every kernel that maps one thread to one element needs a bounds check.")
    add_code(doc, r"""
        __global__ void saxpy(float a, const float* x, float* y, int n) {
            int i = blockIdx.x * blockDim.x + threadIdx.x;
            if (i < n) y[i] = a * x[i] + y[i];
        }
    """)
    add_bullets(doc, [
        "Always check cudaGetLastError after a kernel launch while learning.",
        "Use cudaDeviceSynchronize when you need to catch asynchronous kernel errors immediately.",
        "Test small sizes, odd sizes, and sizes smaller than one block.",
        "Compare GPU output against a simple CPU reference before optimizing.",
    ])

    doc.add_heading("5. Grid-Stride Loops", level=1)
    doc.add_paragraph("A grid-stride loop lets each thread process multiple elements. This makes kernels work for arbitrary N and lets you tune launch size independently from input size.")
    add_code(doc, r"""
        __global__ void add(float* x, int n) {
            for (int i = blockIdx.x * blockDim.x + threadIdx.x;
                 i < n;
                 i += blockDim.x * gridDim.x) {
                x[i] += 1.0f;
            }
        }
    """)

    doc.add_heading("6. Shared Memory And Tiling", level=1)
    doc.add_paragraph("Shared memory is a programmer-managed cache scoped to a block. In matrix multiplication, tiling loads submatrices into shared memory so many threads reuse the same global-memory values.")
    add_bullets(doc, [
        "Each block cooperatively loads a tile from A and a tile from B.",
        "Threads synchronize after loading tiles.",
        "Each thread accumulates one output element using tile data.",
        "The pattern reduces global reads from repeated per-thread loads to reused block-level loads.",
    ])
    add_code(doc, r"""
        __shared__ float As[TILE][TILE];
        __shared__ float Bs[TILE][TILE];
        As[threadIdx.y][threadIdx.x] = A[row * N + tiled_col];
        Bs[threadIdx.y][threadIdx.x] = B[tiled_row * N + col];
        __syncthreads();
    """)

    doc.add_heading("7. Synchronization", level=1)
    add_table(doc, ["Tool", "Scope", "Use"], [
        ["__syncthreads()", "One block", "Barrier before consuming shared data."],
        ["Atomic operations", "Address/global or shared", "Correctness for concurrent updates; can serialize."],
        ["cudaDeviceSynchronize()", "Whole device from host", "Wait for queued work before CPU reads or exits."],
        ["Stream sync", "One stream", "Wait for ordered work in a stream."],
        ["Events", "Device timestamp/order", "Measure GPU time or create cross-stream dependencies."],
    ], widths=[1.8, 1.8, 2.9])
    callout(doc, "Common trap", "__syncthreads() only synchronizes threads in the same block. CUDA has no cheap global barrier inside a normal kernel.", "FEE2E2")

    doc.add_heading("8. Performance Model", level=1)
    doc.add_picture(str(FIGS / "sm_model.png"), width=Inches(6.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("A CUDA performance explanation usually starts with three questions: Is the kernel memory-bound or compute-bound? Is memory access coalesced? Is there enough occupancy and instruction-level work to hide latency?")
    add_table(doc, ["Concept", "Plain-English meaning", "How to improve"], [
        ["Occupancy", "How many warps can reside on an SM.", "Reduce excessive registers/shared memory; choose reasonable block sizes."],
        ["Latency hiding", "Switch to another ready warp while one waits.", "Expose enough parallelism and independent work."],
        ["Arithmetic intensity", "Math per byte moved.", "Reuse data through shared memory, caches, or better algorithms."],
        ["Divergence", "Threads in a warp take different branches.", "Group similar work; simplify branch-heavy kernels."],
        ["Coalescing", "Warp memory addresses form efficient transactions.", "Use structure-of-arrays and contiguous indexing."],
        ["Bank conflicts", "Shared memory accesses collide by bank.", "Pad or reorganize shared arrays."],
    ], widths=[1.4, 2.55, 2.5])

    doc.add_heading("9. Profiling Workflow", level=1)
    doc.add_picture(str(FIGS / "optimization_loop.png"), width=Inches(6.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_bullets(doc, [
        "Use CPU timers only for end-to-end user-visible latency.",
        "Use CUDA events for GPU elapsed time between operations.",
        "Use Nsight Systems to see CPU/GPU overlap, transfers, gaps, and stream timelines.",
        "Use Nsight Compute to inspect one kernel: occupancy, memory throughput, cache behavior, divergence, and instruction mix.",
        "Change one thing at a time; keep the CPU reference and correctness tests alive.",
    ])
    add_code(doc, r"""
        cudaEvent_t start, stop;
        cudaEventCreate(&start);
        cudaEventCreate(&stop);
        cudaEventRecord(start);
        kernel<<<blocks, threads>>>(...);
        cudaEventRecord(stop);
        cudaEventSynchronize(stop);
        float ms = 0.0f;
        cudaEventElapsedTime(&ms, start, stop);
    """)

    doc.add_heading("10. Streams, Async Copies, And Overlap", level=1)
    doc.add_paragraph("A CUDA stream is an ordered queue of work. Operations in the same stream run in issue order. Independent streams may overlap if the hardware and dependencies allow it.")
    add_bullets(doc, [
        "Use cudaMemcpyAsync plus pinned host memory for asynchronous host-device transfers.",
        "Use multiple streams to pipeline copy, compute, and copy-back chunks.",
        "Avoid synchronizing the whole device unless necessary.",
        "Remember that default-stream behavior matters; know whether your code uses legacy or per-thread default stream semantics.",
    ])

    doc.add_heading("11. AI/ML CUDA Concepts", level=1)
    add_table(doc, ["Topic", "What to know for interviews"], [
        ["Tensor cores", "Specialized matrix-multiply hardware used by deep learning kernels."],
        ["Mixed precision", "FP16/BF16/TF32/FP8 trade precision for throughput and memory savings."],
        ["cuBLAS/cuDNN", "Use vendor libraries for GEMM, convolutions, reductions, and neural-network primitives."],
        ["Custom ops", "Write CUDA extensions when framework primitives do not express your operation efficiently."],
        ["Memory bandwidth", "Large models are often limited by data movement, not raw FLOPs."],
        ["Batching", "Improves throughput by increasing parallel work and matrix sizes."],
        ["Kernel fusion", "Combines operations to reduce intermediate memory traffic and launch overhead."],
    ], widths=[1.55, 4.9])
    callout(doc, "Domain-change advice", "For CUDA AI/ML roles, become fluent in C++, linear algebra, GPU memory behavior, profiling, and at least one ML runtime such as PyTorch extensions, TensorRT, or Triton.", "DCFCE7")

    doc.add_heading("12. Debugging And Correctness", level=1)
    add_bullets(doc, [
        "Compile debug builds with -G while debugging, then remove -G for performance measurements.",
        "Use compute-sanitizer to catch illegal memory access, race conditions, and uninitialized values.",
        "Keep CPU reference implementations for kernels with tricky indexing.",
        "Check every CUDA API call with a macro during development.",
        "Treat printf inside kernels as a debugging tool only; it changes timing and can be buffered.",
    ])
    add_code(doc, r"""
        nvcc -G 02_vector_add.cu -o debug_vector_add.exe
        compute-sanitizer .\debug_vector_add.exe
    """)

    doc.add_heading("13. Runnable Sample Programs", level=1)
    add_table(doc, ["File", "Concept", "Compile command"], [
        ["01_hello.cu", "Kernel launch and thread coordinates", "nvcc 01_hello.cu -o 01_hello.exe"],
        ["02_vector_add.cu", "Memory allocation, copies, error checking", "nvcc 02_vector_add.cu -o 02_vector_add.exe"],
        ["03_matrix_mul_tiled.cu", "Shared-memory tiling", "nvcc 03_matrix_mul_tiled.cu -o 03_matrix_mul_tiled.exe"],
        ["04_reduction.cu", "Parallel reduction", "nvcc 04_reduction.cu -o 04_reduction.exe"],
        ["05_streams_async.cu", "Pinned memory and async stream work", "nvcc 05_streams_async.cu -o 05_streams_async.exe"],
        ["06_unified_memory.cu", "Managed memory", "nvcc 06_unified_memory.cu -o 06_unified_memory.exe"],
    ], widths=[1.55, 2.4, 2.5])

    doc.add_heading("Vector Add Example", level=2)
    add_code(doc, SAMPLE_FILES["02_vector_add.cu"])
    doc.add_heading("Tiled Matrix Multiplication Example", level=2)
    add_code(doc, SAMPLE_FILES["03_matrix_mul_tiled.cu"])
    doc.add_heading("Reduction Example", level=2)
    add_code(doc, SAMPLE_FILES["04_reduction.cu"])

    doc.add_heading("14. Interview Questions And Model Answers", level=1)
    qa = [
        ("What happens when you launch kernel<<<blocks, threads>>>?", "The host queues work for the GPU. The launch creates a grid of blocks; each block contains threads. The launch is usually asynchronous with respect to the CPU."),
        ("Why do we need an if (i < n) guard?", "Grid sizes are commonly rounded up. Extra threads would otherwise access memory past the valid range."),
        ("What is a warp?", "A hardware scheduling group of threads, commonly 32. Threads in a warp execute together, so divergent branches serialize paths."),
        ("Shared memory vs global memory?", "Global memory is large and high latency. Shared memory is per-block, much lower latency, and used for cooperation and data reuse."),
        ("What is occupancy?", "The ratio of active warps on an SM to the hardware maximum. It matters because more resident warps can hide latency, but maximum occupancy is not always maximum performance."),
        ("How do you optimize a memory-bound kernel?", "Reduce bytes moved, coalesce accesses, reuse data through shared memory/cache, avoid unnecessary transfers, and consider fusion."),
        ("When are atomics appropriate?", "When multiple threads update the same location and exact correctness matters. Use them carefully because contention can serialize execution."),
        ("How do streams improve performance?", "They let independent copies and kernels be queued separately so the GPU can overlap transfers and compute when dependencies and hardware permit."),
        ("How do you profile CUDA?", "Start with end-to-end timing, use CUDA events for GPU work, Nsight Systems for timeline gaps/overlap, and Nsight Compute for kernel metrics."),
        ("Why use cuBLAS instead of writing GEMM?", "Vendor libraries are heavily optimized for architecture-specific tiling, tensor cores, memory layouts, and edge cases."),
    ]
    for q, a in qa:
        p = doc.add_paragraph()
        p.add_run("Q: " + q).bold = True
        doc.add_paragraph("A: " + a)

    doc.add_heading("15. 0-to-10 Roadmap", level=1)
    add_table(doc, ["Level", "Focus", "You can prove it by"], [
        ["0", "C++ refresh and toolchain", "Compile C++ and CUDA hello-world programs."],
        ["1", "Kernel launches and indexing", "Write vector add and SAXPY from memory."],
        ["2", "CUDA memory APIs", "Explain malloc/copy/free and error checking."],
        ["3", "Thread/block/grid mapping", "Handle 1D and 2D data correctly."],
        ["4", "Shared memory", "Implement tiled matrix multiplication."],
        ["5", "Reductions and atomics", "Implement sum reduction and explain contention."],
        ["6", "Streams and events", "Overlap copy/compute and time kernels correctly."],
        ["7", "Profiling", "Use Nsight to justify a bottleneck."],
        ["8", "Advanced memory/performance", "Discuss coalescing, bank conflicts, occupancy, divergence."],
        ["9", "AI/ML systems", "Use cuBLAS/cuDNN or write a PyTorch CUDA extension."],
        ["10", "Expert interview readiness", "Whiteboard kernels, debug failures, and defend optimization choices."],
    ], widths=[0.55, 2.4, 3.5])

    doc.add_heading("16. Sources And Further Reading", level=1)
    add_bullets(doc, [
        "NVIDIA CUDA C++ Programming Guide: https://docs.nvidia.com/cuda/cuda-c-programming-guide/",
        "NVIDIA CUDA C++ Best Practices Guide: https://docs.nvidia.com/cuda/cuda-c-best-practices-guide/",
        "NVIDIA CUDA Runtime API: https://docs.nvidia.com/cuda/cuda-runtime-api/",
        "NVIDIA CUDA Installation Guide for Microsoft Windows: https://docs.nvidia.com/cuda/cuda-installation-guide-microsoft-windows/",
        "NVIDIA Nsight Compute User Guide: https://docs.nvidia.com/nsight-compute/",
        "NVIDIA Deep Learning Performance Documentation: https://docs.nvidia.com/deeplearning/performance/",
    ])

    doc.save(OUT)


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="PdfTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=24,
        leading=28,
        textColor=colors.HexColor("#245B91"),
        alignment=TA_CENTER,
        spaceAfter=10,
    ))
    styles.add(ParagraphStyle(
        name="PdfSubtitle",
        parent=styles["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=10,
        leading=14,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#334155"),
        spaceAfter=14,
    ))
    styles.add(ParagraphStyle(
        name="H1x",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=20,
        textColor=colors.HexColor("#245B91"),
        spaceBefore=14,
        spaceAfter=7,
    ))
    styles.add(ParagraphStyle(
        name="H2x",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12.5,
        leading=16,
        textColor=colors.HexColor("#1F7A72"),
        spaceBefore=10,
        spaceAfter=5,
    ))
    styles.add(ParagraphStyle(
        name="BodyX",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=13,
        spaceAfter=6,
        textColor=colors.HexColor("#172033"),
    ))
    styles.add(ParagraphStyle(
        name="CodeX",
        parent=styles["Code"],
        fontName="Courier",
        fontSize=6.7,
        leading=8.2,
        leftIndent=6,
        rightIndent=6,
        borderPadding=5,
        borderColor=colors.HexColor("#CBD5E1"),
        borderWidth=0.5,
        backColor=colors.HexColor("#F8FAFC"),
        spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        name="SmallX",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=8,
        leading=10,
        textColor=colors.HexColor("#334155"),
    ))
    return styles


def ptable(headers, rows, widths):
    data = [[Paragraph(f"<b>{h}</b>", PDF_STYLES["SmallX"]) for h in headers]]
    for row in rows:
        data.append([Paragraph(str(c), PDF_STYLES["SmallX"]) for c in row])
    table = Table(data, colWidths=[w * inch for w in widths], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
        ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#172033")),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#CBD5E1")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
    ]))
    return [table, Spacer(1, 8)]


def plist(items):
    return ListFlowable(
        [ListItem(Paragraph(item, PDF_STYLES["BodyX"])) for item in items],
        bulletType="bullet",
        start="circle",
        leftIndent=16,
    )


def pnum(items):
    return ListFlowable(
        [ListItem(Paragraph(item, PDF_STYLES["BodyX"])) for item in items],
        bulletType="1",
        leftIndent=18,
    )


def codeblock(code):
    return Preformatted(dedent(code).strip(), PDF_STYLES["CodeX"], maxLineLength=95)


def pdf_img(name):
    return PdfImage(str(FIGS / name), width=6.6 * inch, height=None)


def add_page_number(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#64748B"))
    canvas.drawCentredString(4.25 * inch, 0.38 * inch, f"CUDA C++ AI/ML Interview Handbook | Page {doc.page}")
    canvas.restoreState()


def build_pdf():
    make_figures()
    write_samples()
    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=letter,
        rightMargin=0.65 * inch,
        leftMargin=0.65 * inch,
        topMargin=0.62 * inch,
        bottomMargin=0.62 * inch,
    )
    story = []
    S = PDF_STYLES
    story.append(Paragraph("CUDA C++ for AI/ML Interviews", S["PdfTitle"]))
    story.append(Paragraph("Beginner to expert roadmap, architecture notes, runnable code, execution commands, and interview drills", S["PdfSubtitle"]))
    story.append(Paragraph("Audience: a full-stack engineer moving into C++ CUDA, GPU computing, AI/ML systems, inference, training, and performance engineering.", S["BodyX"]))
    story.append(Paragraph("<b>How to use this PDF:</b> read it in passes: first understand the mental model, then compile every sample, then profile and explain each optimization aloud.", S["BodyX"]))

    story += [Paragraph("0. What CUDA C++ Is", S["H1x"])]
    story.append(Paragraph("CUDA C++ is C++ plus NVIDIA GPU extensions. You write normal host code that runs on the CPU and kernel code that runs on the GPU. The GPU is a throughput machine designed to run many lightweight threads and hide memory latency with massive concurrency.", S["BodyX"]))
    story.append(plist([
        "CPU code is host code. GPU code is device code.",
        "A kernel is a GPU function launched from the CPU with <<<grid, block>>> syntax.",
        "A warp is a hardware scheduling group, commonly 32 threads.",
        "A block is a group of threads that can synchronize and share fast shared memory.",
        "A grid is all blocks created by one kernel launch.",
    ]))
    story.append(pdf_img("execution_hierarchy.png"))

    story += [Paragraph("1. Setup And First Execution", S["H1x"])]
    story.append(pnum([
        "Install an NVIDIA driver that supports your GPU.",
        "Install the CUDA Toolkit. On Windows, install a compatible Visual Studio C++ toolchain first.",
        "Confirm the toolchain with: nvcc --version",
        "Confirm the GPU and driver with: nvidia-smi",
        "Compile a .cu file with: nvcc file.cu -o file.exe",
        "Run the executable from PowerShell: .\\file.exe",
    ]))
    story.append(codeblock(r"""
        cd C:\Users\ygirish\Documents\Codex\2026-05-28\hi\cuda_samples
        nvcc 01_hello.cu -o 01_hello.exe
        .\01_hello.exe
    """))

    story += [Paragraph("2. CUDA Language Keywords", S["H1x"])]
    story += ptable(["Keyword", "Meaning", "Interview signal"], [
        ["__global__", "Function runs on GPU and is launched by CPU.", "Used for kernels; returns void."],
        ["__device__", "Function or variable lives on GPU device.", "Callable from device code."],
        ["__host__", "Function runs on CPU host.", "Can combine with __device__ for dual-use helpers."],
        ["__shared__", "Per-block shared memory.", "Fast cooperation inside one block."],
        ["__syncthreads()", "Barrier for threads in a block.", "Use before reading shared data written by peers."],
        ["threadIdx/blockIdx", "Built-in coordinates.", "Used to map work to data."],
        ["blockDim/gridDim", "Launch dimensions.", "Used for bounds and grid-stride loops."],
    ], [1.2, 2.8, 2.3])

    story += [Paragraph("3. Memory Model", S["H1x"]), pdf_img("memory_hierarchy.png")]
    story += ptable(["Memory", "Lifetime / scope", "Typical use", "Pitfall"], [
        ["Registers", "One thread", "Local scalars", "Too many registers can reduce occupancy."],
        ["Local memory", "One thread, backed by global", "Spilled arrays/variables", "Slow if compiler spills registers."],
        ["Shared memory", "One block", "Tiling, reductions, stencil reuse", "Bank conflicts and missing barriers."],
        ["Global memory", "Whole device allocation", "Large arrays and tensors", "Latency, non-coalesced access."],
        ["Constant memory", "Device-wide read-only", "Small broadcast constants", "Bad for random per-thread access."],
        ["Unified memory", "Managed CPU/GPU address", "Simplicity and prototyping", "Page migration surprises."],
        ["Pinned host memory", "CPU page-locked", "Async copies and overlap", "Overuse hurts system memory."],
    ], [1.1, 1.3, 1.9, 1.9])
    story.append(Paragraph("<b>Interview answer:</b> coalescing means adjacent threads in a warp access adjacent memory addresses, allowing the hardware to combine requests efficiently.", S["BodyX"]))

    story += [Paragraph("4. Writing Correct Kernels", S["H1x"])]
    story.append(Paragraph("The most common CUDA beginner bug is forgetting that you usually launch more threads than data elements. Every kernel that maps one thread to one element needs a bounds check.", S["BodyX"]))
    story.append(codeblock(r"""
        __global__ void saxpy(float a, const float* x, float* y, int n) {
            int i = blockIdx.x * blockDim.x + threadIdx.x;
            if (i < n) y[i] = a * x[i] + y[i];
        }
    """))
    story.append(plist([
        "Always check cudaGetLastError after a kernel launch while learning.",
        "Use cudaDeviceSynchronize when you need to catch asynchronous kernel errors immediately.",
        "Test small sizes, odd sizes, and sizes smaller than one block.",
        "Compare GPU output against a simple CPU reference before optimizing.",
    ]))

    story += [Paragraph("5. Grid-Stride Loops", S["H1x"])]
    story.append(Paragraph("A grid-stride loop lets each thread process multiple elements. This makes kernels work for arbitrary N and lets you tune launch size independently from input size.", S["BodyX"]))
    story.append(codeblock(r"""
        __global__ void add(float* x, int n) {
            for (int i = blockIdx.x * blockDim.x + threadIdx.x;
                 i < n;
                 i += blockDim.x * gridDim.x) {
                x[i] += 1.0f;
            }
        }
    """))

    story += [Paragraph("6. Shared Memory And Tiling", S["H1x"])]
    story.append(Paragraph("Shared memory is a programmer-managed cache scoped to a block. In matrix multiplication, tiling loads submatrices into shared memory so many threads reuse the same global-memory values.", S["BodyX"]))
    story.append(plist([
        "Each block cooperatively loads a tile from A and a tile from B.",
        "Threads synchronize after loading tiles.",
        "Each thread accumulates one output element using tile data.",
        "The pattern reduces global reads from repeated per-thread loads to reused block-level loads.",
    ]))
    story.append(codeblock(r"""
        __shared__ float As[TILE][TILE];
        __shared__ float Bs[TILE][TILE];
        As[threadIdx.y][threadIdx.x] = A[row * N + tiled_col];
        Bs[threadIdx.y][threadIdx.x] = B[tiled_row * N + col];
        __syncthreads();
    """))

    story += [Paragraph("7. Synchronization", S["H1x"])]
    story += ptable(["Tool", "Scope", "Use"], [
        ["__syncthreads()", "One block", "Barrier before consuming shared data."],
        ["Atomic operations", "Address/global or shared", "Correctness for concurrent updates; can serialize."],
        ["cudaDeviceSynchronize()", "Whole device from host", "Wait for queued work before CPU reads or exits."],
        ["Stream sync", "One stream", "Wait for ordered work in a stream."],
        ["Events", "Device timestamp/order", "Measure GPU time or create cross-stream dependencies."],
    ], [1.7, 1.7, 2.9])
    story.append(Paragraph("<b>Common trap:</b> __syncthreads() only synchronizes threads in the same block. CUDA has no cheap global barrier inside a normal kernel.", S["BodyX"]))

    story += [PageBreak(), Paragraph("8. Performance Model", S["H1x"]), pdf_img("sm_model.png")]
    story.append(Paragraph("A CUDA performance explanation usually starts with three questions: Is the kernel memory-bound or compute-bound? Is memory access coalesced? Is there enough occupancy and instruction-level work to hide latency?", S["BodyX"]))
    story += ptable(["Concept", "Plain-English meaning", "How to improve"], [
        ["Occupancy", "How many warps can reside on an SM.", "Reduce excessive registers/shared memory; choose reasonable block sizes."],
        ["Latency hiding", "Switch to another ready warp while one waits.", "Expose enough parallelism and independent work."],
        ["Arithmetic intensity", "Math per byte moved.", "Reuse data through shared memory, caches, or better algorithms."],
        ["Divergence", "Threads in a warp take different branches.", "Group similar work; simplify branch-heavy kernels."],
        ["Coalescing", "Warp memory addresses form efficient transactions.", "Use structure-of-arrays and contiguous indexing."],
        ["Bank conflicts", "Shared memory accesses collide by bank.", "Pad or reorganize shared arrays."],
    ], [1.3, 2.5, 2.5])

    story += [Paragraph("9. Profiling Workflow", S["H1x"]), pdf_img("optimization_loop.png")]
    story.append(plist([
        "Use CPU timers only for end-to-end user-visible latency.",
        "Use CUDA events for GPU elapsed time between operations.",
        "Use Nsight Systems to see CPU/GPU overlap, transfers, gaps, and stream timelines.",
        "Use Nsight Compute to inspect one kernel: occupancy, memory throughput, cache behavior, divergence, and instruction mix.",
        "Change one thing at a time; keep the CPU reference and correctness tests alive.",
    ]))
    story.append(codeblock(r"""
        cudaEvent_t start, stop;
        cudaEventCreate(&start);
        cudaEventCreate(&stop);
        cudaEventRecord(start);
        kernel<<<blocks, threads>>>(...);
        cudaEventRecord(stop);
        cudaEventSynchronize(stop);
        float ms = 0.0f;
        cudaEventElapsedTime(&ms, start, stop);
    """))

    story += [Paragraph("10. Streams, Async Copies, And Overlap", S["H1x"])]
    story.append(Paragraph("A CUDA stream is an ordered queue of work. Operations in the same stream run in issue order. Independent streams may overlap if the hardware and dependencies allow it.", S["BodyX"]))
    story.append(plist([
        "Use cudaMemcpyAsync plus pinned host memory for asynchronous host-device transfers.",
        "Use multiple streams to pipeline copy, compute, and copy-back chunks.",
        "Avoid synchronizing the whole device unless necessary.",
        "Remember that default-stream behavior matters; know whether your code uses legacy or per-thread default stream semantics.",
    ]))

    story += [Paragraph("11. AI/ML CUDA Concepts", S["H1x"])]
    story += ptable(["Topic", "What to know for interviews"], [
        ["Tensor cores", "Specialized matrix-multiply hardware used by deep learning kernels."],
        ["Mixed precision", "FP16/BF16/TF32/FP8 trade precision for throughput and memory savings."],
        ["cuBLAS/cuDNN", "Use vendor libraries for GEMM, convolutions, reductions, and neural-network primitives."],
        ["Custom ops", "Write CUDA extensions when framework primitives do not express your operation efficiently."],
        ["Memory bandwidth", "Large models are often limited by data movement, not raw FLOPs."],
        ["Batching", "Improves throughput by increasing parallel work and matrix sizes."],
        ["Kernel fusion", "Combines operations to reduce intermediate memory traffic and launch overhead."],
    ], [1.5, 4.8])

    story += [Paragraph("12. Debugging And Correctness", S["H1x"])]
    story.append(plist([
        "Compile debug builds with -G while debugging, then remove -G for performance measurements.",
        "Use compute-sanitizer to catch illegal memory access, race conditions, and uninitialized values.",
        "Keep CPU reference implementations for kernels with tricky indexing.",
        "Check every CUDA API call with a macro during development.",
        "Treat printf inside kernels as a debugging tool only; it changes timing and can be buffered.",
    ]))
    story.append(codeblock(r"""
        nvcc -G 02_vector_add.cu -o debug_vector_add.exe
        compute-sanitizer .\debug_vector_add.exe
    """))

    story += [PageBreak(), Paragraph("13. Runnable Sample Programs", S["H1x"])]
    story += ptable(["File", "Concept", "Compile command"], [
        ["01_hello.cu", "Kernel launch and thread coordinates", "nvcc 01_hello.cu -o 01_hello.exe"],
        ["02_vector_add.cu", "Memory allocation, copies, error checking", "nvcc 02_vector_add.cu -o 02_vector_add.exe"],
        ["03_matrix_mul_tiled.cu", "Shared-memory tiling", "nvcc 03_matrix_mul_tiled.cu -o 03_matrix_mul_tiled.exe"],
        ["04_reduction.cu", "Parallel reduction", "nvcc 04_reduction.cu -o 04_reduction.exe"],
        ["05_streams_async.cu", "Pinned memory and async stream work", "nvcc 05_streams_async.cu -o 05_streams_async.exe"],
        ["06_unified_memory.cu", "Managed memory", "nvcc 06_unified_memory.cu -o 06_unified_memory.exe"],
    ], [1.4, 2.3, 2.6])

    for heading, fname in [
        ("Vector Add Example", "02_vector_add.cu"),
        ("Tiled Matrix Multiplication Example", "03_matrix_mul_tiled.cu"),
        ("Reduction Example", "04_reduction.cu"),
    ]:
        story.append(Paragraph(heading, S["H2x"]))
        story.append(codeblock((SAMPLES / fname).read_text(encoding="utf-8")))

    story += [Paragraph("14. Interview Questions And Model Answers", S["H1x"])]
    qa = [
        ("What happens when you launch kernel<<<blocks, threads>>>?", "The host queues work for the GPU. The launch creates a grid of blocks; each block contains threads. The launch is usually asynchronous with respect to the CPU."),
        ("Why do we need an if (i < n) guard?", "Grid sizes are commonly rounded up. Extra threads would otherwise access memory past the valid range."),
        ("What is a warp?", "A hardware scheduling group of threads, commonly 32. Threads in a warp execute together, so divergent branches serialize paths."),
        ("Shared memory vs global memory?", "Global memory is large and high latency. Shared memory is per-block, much lower latency, and used for cooperation and data reuse."),
        ("What is occupancy?", "The ratio of active warps on an SM to the hardware maximum. It matters because more resident warps can hide latency, but maximum occupancy is not always maximum performance."),
        ("How do you optimize a memory-bound kernel?", "Reduce bytes moved, coalesce accesses, reuse data through shared memory/cache, avoid unnecessary transfers, and consider fusion."),
        ("When are atomics appropriate?", "When multiple threads update the same location and exact correctness matters. Use them carefully because contention can serialize execution."),
        ("How do streams improve performance?", "They let independent copies and kernels be queued separately so the GPU can overlap transfers and compute when dependencies and hardware permit."),
        ("How do you profile CUDA?", "Start with end-to-end timing, use CUDA events for GPU work, Nsight Systems for timeline gaps/overlap, and Nsight Compute for kernel metrics."),
        ("Why use cuBLAS instead of writing GEMM?", "Vendor libraries are heavily optimized for architecture-specific tiling, tensor cores, memory layouts, and edge cases."),
    ]
    for q, a in qa:
        story.append(Paragraph(f"<b>Q: {q}</b>", S["BodyX"]))
        story.append(Paragraph(f"A: {a}", S["BodyX"]))

    story += [Paragraph("15. 0-to-10 Roadmap", S["H1x"])]
    story += ptable(["Level", "Focus", "You can prove it by"], [
        ["0", "C++ refresh and toolchain", "Compile C++ and CUDA hello-world programs."],
        ["1", "Kernel launches and indexing", "Write vector add and SAXPY from memory."],
        ["2", "CUDA memory APIs", "Explain malloc/copy/free and error checking."],
        ["3", "Thread/block/grid mapping", "Handle 1D and 2D data correctly."],
        ["4", "Shared memory", "Implement tiled matrix multiplication."],
        ["5", "Reductions and atomics", "Implement sum reduction and explain contention."],
        ["6", "Streams and events", "Overlap copy/compute and time kernels correctly."],
        ["7", "Profiling", "Use Nsight to justify a bottleneck."],
        ["8", "Advanced memory/performance", "Discuss coalescing, bank conflicts, occupancy, divergence."],
        ["9", "AI/ML systems", "Use cuBLAS/cuDNN or write a PyTorch CUDA extension."],
        ["10", "Expert interview readiness", "Whiteboard kernels, debug failures, and defend optimization choices."],
    ], [0.55, 2.3, 3.45])

    story += [Paragraph("16. Sources And Further Reading", S["H1x"])]
    story.append(plist([
        "NVIDIA CUDA C++ Programming Guide: https://docs.nvidia.com/cuda/cuda-c-programming-guide/",
        "NVIDIA CUDA C++ Best Practices Guide: https://docs.nvidia.com/cuda/cuda-c-best-practices-guide/",
        "NVIDIA CUDA Runtime API: https://docs.nvidia.com/cuda/cuda-runtime-api/",
        "NVIDIA CUDA Installation Guide for Microsoft Windows: https://docs.nvidia.com/cuda/cuda-installation-guide-microsoft-windows/",
        "NVIDIA Nsight Compute User Guide: https://docs.nvidia.com/nsight-compute/",
        "NVIDIA Deep Learning Performance Documentation: https://docs.nvidia.com/deeplearning/performance/",
    ]))

    doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)


if __name__ == "__main__":
    PDF_STYLES = pdf_styles()
    build_doc()
    build_pdf()
    print(OUT)
    print(PDF_OUT)
